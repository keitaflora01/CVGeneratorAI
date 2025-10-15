from datetime import datetime
import logging
import os
from django.template.loader import render_to_string
import re
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.contrib.auth import get_user_model
import json
import time
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document as DocxDocument
import markdown
from weasyprint import HTML
import google.generativeai as genai
from decouple import config, UndefinedValueError
from tavily import TavilyClient
from Agent.models import Document, EtapeTraitement, CVImage

# Configure logging
logger = logging.getLogger(__name__)

User = get_user_model()

# Configure API clients safely
try:
    GEMINI_API_KEY = config('GEMINI_API_KEY')
    TAVILY_API_KEY = config('TAVILY_API_KEY')
    genai.configure(api_key=GEMINI_API_KEY)
    tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
except UndefinedValueError as e:
    logger.error(f"Missing environment variable: {str(e)}")
    raise Exception(f"Missing environment variable: {str(e)}")

def get_available_templates():
    """Dynamically retrieve available CV and LM templates from directories."""
    templates = []
    base_dir = settings.TEMPLATES[0]['DIRS'][0]  # Assumes templates are in the first template directory

    # CV templates
    cv_template_dir = os.path.join(base_dir, 'cv-templates')
    if os.path.exists(cv_template_dir):
        for filename in os.listdir(cv_template_dir):
            if filename.endswith('.html'):
                template_id = filename.replace('.html', '')
                template_name = template_id.replace('-', ' ').title()
                templates.append({
                    'id': template_id,
                    'type': 'CV',
                    'name': template_name,
                    'path': os.path.join('cv-templates', filename)
                })

    # LM templates
    lm_template_dir = os.path.join(base_dir, 'lettres-motivation')
    if os.path.exists(lm_template_dir):
        for filename in os.listdir(lm_template_dir):
            if filename.endswith('.html'):
                template_id = filename.replace('.html', '')
                template_name = template_id.replace('-', ' ').title()
                templates.append({
                    'id': template_id,
                    'type': 'LM',
                    'name': template_name,
                    'path': os.path.join('lettres-motivation', filename)
                })

    logger.debug(f"Retrieved templates: {[t['name'] for t in templates]}")
    return templates

@csrf_exempt
def generate_document(request):
    """Generate CV or Letter of Motivation using Gemini and Tavily APIs"""
    logger.info(f"[generate_document] {time.strftime('%Y-%m-%d %H:%M:%S')} | Method: {request.method} | Path: {request.path} | Headers: {dict(request.headers)}")

    if request.method == 'POST':
        try:
            logger.debug(f"POST Data: {request.POST.dict()}")
            logger.debug(f"Files: {request.FILES}")

            # Extract form data
            target_role = request.POST.get('targetRole', '').strip()
            company = request.POST.get('company', '').strip()
            keywords = request.POST.get('keywords', '').strip()
            tone = request.POST.get('tone', 'professionnel')
            job_description = request.POST.get('jobDescription', '').strip()
            document_type = request.POST.get('documentType', 'CV')
            linkedin_url = request.POST.get('linkedin_url', '').strip()
            github_url = request.POST.get('github_url', '').strip()
            telephone = request.POST.get('telephone', '').strip()
            langue = request.POST.get('langue', 'fr')
            template_id = request.POST.get('template_id', 'template1-moderne-bleu')  # Default
            template_utilise = request.POST.get('template_utilise', 'Template1 Moderne Bleu')
            skills = [skill.strip() for skill in request.POST.get('skills', '').split(',') if skill.strip()]
            try:
                experiences = json.loads(request.POST.get('experiences', '[]'))
                education = json.loads(request.POST.get('education', '[]'))
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error for experiences/education: {str(e)}")
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': False,
                        'error': 'Format invalide pour les expériences ou l\'éducation'
                    })
                return render(request, 'user/generate.html', {
                    'error': 'Format invalide pour les expériences ou l\'éducation',
                    'templates': get_available_templates()
                })

            if not target_role or not job_description:
                logger.warning("Missing required fields: targetRole or jobDescription")
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': False,
                        'error': 'Le poste ciblé et la description sont obligatoires'
                    })
                return render(request, 'user/generate.html', {
                    'error': 'Le poste ciblé et la description sont obligatoires',
                    'templates': get_available_templates()
                })

            # Create or update document record
            user = request.user if request.user.is_authenticated else None
            logger.info(f"Creating document for user: {user.email if user else 'Anonymous'}")
            doc_id = request.POST.get('doc_id')
            if doc_id:
                document = get_object_or_404(Document, id=doc_id, user=user)
                document.type = document_type
                document.titre = f"{document_type} for {target_role}"
                document.poste = target_role
                document.entreprise = company
                document.linkedin_url = linkedin_url
                document.github_url = github_url
                document.telephone = telephone
                document.langue = langue
                document.template_utilise = template_utilise
                document.statut = 'processing'
                document.metadata = {
                    'keywords': keywords,
                    'tone': tone,
                    'job_description_preview': job_description[:100] + '...' if job_description else '',
                    'linkedin_url': linkedin_url,
                    'github_url': github_url,
                    'telephone': telephone,
                    'langue': langue,
                    'template_id': template_id,
                    'template_utilise': template_utilise
                }
            else:
                document = Document.objects.create(
                    user=user,
                    type=document_type,
                    titre=f"{document_type} for {target_role}",
                    poste=target_role,
                    entreprise=company,
                    linkedin_url=linkedin_url,
                    github_url=github_url,
                    telephone=telephone,
                    langue=langue,
                    template_utilise=template_utilise,
                    statut='processing',
                    metadata={
                        'keywords': keywords,
                        'tone': tone,
                        'job_description_preview': job_description[:100] + '...' if job_description else '',
                        'linkedin_url': linkedin_url,
                        'github_url': github_url,
                        'telephone': telephone,
                        'langue': langue,
                        'template_id': template_id,
                        'template_utilise': template_utilise
                    }
                )
            logger.info(f"Document {'updated' if doc_id else 'created'} with ID: {document.id}")

            # Handle CV image with enhanced processing
            if 'cv_image' in request.FILES and document_type == 'CV':
                cv_image = request.FILES['cv_image']
                logger.info(f"Processing CV image: {cv_image.name}, size: {cv_image.size} bytes")
                
                # Enhanced image validation
                if cv_image.size > 2 * 1024 * 1024:  # 2MB limit
                    logger.warning(f"Image size exceeds 2MB: {cv_image.size}")
                    document.delete()
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return JsonResponse({
                            'success': False,
                            'error': 'L\'image ne doit pas dépasser 2MB. Veuillez compresser votre image.'
                        })
                    return render(request, 'user/generate.html', {
                        'error': 'L\'image ne doit pas dépasser 2MB. Veuillez compresser votre image.',
                        'templates': get_available_templates()
                    })
                
                # Check file extension
                allowed_extensions = ['.jpg', '.jpeg', '.png', '.gif']
                file_extension = os.path.splitext(cv_image.name)[1].lower()
                if file_extension not in allowed_extensions:
                    logger.warning(f"Invalid image format: {file_extension}")
                    document.delete()
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return JsonResponse({
                            'success': False,
                            'error': 'Format d\'image non supporté. Utilisez JPG, PNG ou GIF.'
                        })
                    return render(request, 'user/generate.html', {
                        'error': 'Format d\'image non supporté. Utilisez JPG, PNG ou GIF.',
                        'templates': get_available_templates()
                    })
                
                try:
                    # Create or update CV image
                    cv_image_obj, created = CVImage.objects.get_or_create(
                        document=document,
                        defaults={
                            'image': cv_image,
                            'description': f"Photo professionnelle pour le poste de {target_role}"
                        }
                    )
                    if not created:
                        # Update existing image
                        cv_image_obj.image = cv_image
                        cv_image_obj.description = f"Photo professionnelle pour le poste de {target_role}"
                        cv_image_obj.save()
                    
                    logger.info(f"CV image {'created' if created else 'updated'} successfully for document {document.id}")
                except Exception as e:
                    logger.error(f"Error saving CV image: {str(e)}")
                    document.delete()
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return JsonResponse({
                            'success': False,
                            'error': f'Erreur lors de l\'enregistrement de l\'image: {str(e)}'
                        })
                    return render(request, 'user/generate.html', {
                        'error': f'Erreur lors de l\'enregistrement de l\'image: {str(e)}',
                        'templates': get_available_templates()
                    })

            # Create processing steps
            EtapeTraitement.objects.filter(document=document).delete()
            etapes_data = [
                {"nom": "Job offer analysis", "ordre": 1},
                {"nom": "Profile adaptation", "ordre": 2},
                {"nom": "Content generation", "ordre": 3},
                {"nom": "Optimization", "ordre": 4},
                {"nom": "Final validation", "ordre": 5},
            ]
            for etape_data in etapes_data:
                EtapeTraitement.objects.create(document=document, **etape_data)
            logger.debug("Processing steps created")

            # Fetch additional context using Tavily
            try:
                logger.info(f"Searching Tavily for: {target_role} job requirements {company}")
                tavily_response = tavily_client.search(
                    query=f"{target_role} job requirements {company}",
                    search_depth="basic",
                    max_results=3
                )
                context = "\n".join([result['content'] for result in tavily_response['results']])
                logger.debug(f"Tavily context: {context[:200]}...")
            except Exception as e:
                logger.error(f"Tavily API error: {str(e)}")
                context = "No additional context available."

            # Prepare prompt based on document type and template
            user_data = {
                'name': user.full_name if user and user.is_authenticated and hasattr(user, 'full_name') and user.full_name else user.full_name if user and user.is_authenticated else 'Anonymous',
                'email': user.email if user and user.is_authenticated else 'N/A',
                'linkedin_url': linkedin_url,
                'github_url': github_url,
                'telephone': telephone,
                'skills': skills,
                'experiences': experiences,
                'education': education
            }
            logger.debug(f"User data for prompt: {user_data}")
            prompt = _get_prompt(document_type, target_role, company, keywords, tone, job_description, user_data, context, langue, template_id)
            logger.debug(f"Generated prompt: {prompt[:200]}...")

            # Generate content using Gemini API with enhanced settings
            try:
                logger.info("Calling Gemini API for professional content generation")
                
                # Configure Gemini model with professional settings
                generation_config = {
                    "temperature": 0.8,  # More creative and natural
                    "top_p": 0.9,        # More diverse vocabulary
                    "top_k": 50,         # Wider vocabulary range
                    "max_output_tokens": 3000,  # Sufficient length for professional content
                }
                
                safety_settings = [
                    {
                        "category": "HARM_CATEGORY_HARASSMENT",
                        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                    },
                    {
                        "category": "HARM_CATEGORY_HATE_SPEECH",
                        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                    },
                    {
                        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                    },
                    {
                        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                        "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                    }
                ]
                
                # Try different Gemini models in order of preference
                models_to_try = ['gemini-pro', 'gemini-1.0-pro', 'gemini-2.0-flash']
                generated_content = None
                
                for model_name in models_to_try:
                    try:
                        logger.info(f"Trying Gemini model: {model_name}")
                        model = genai.GenerativeModel(
                            model_name,
                            generation_config=generation_config,
                            safety_settings=safety_settings
                        )
                        
                        logger.debug(f"Sending prompt to Gemini API (length: {len(prompt)} characters)")
                        response = model.generate_content(prompt)
                        
                        if response and response.text:
                            generated_content = response.text.strip()
                            
                            # Clean up the content to remove unwanted formatting
                            if document_type == 'LM':
                                # Remove any remaining bracket structures for motivation letters
                                generated_content = re.sub(r'\[.*?\]', '', generated_content)
                                generated_content = re.sub(r'##.*?\n', '', generated_content)
                                generated_content = re.sub(r'\n\s*\n\s*\n', '\n\n', generated_content)  # Remove excessive line breaks
                                generated_content = generated_content.strip()
                            elif document_type == 'CV':
                                # Remove markdown formatting and clean up CV content
                                generated_content = re.sub(r'\*+', '', generated_content)  # Remove asterisks
                                generated_content = re.sub(r'#+\s*', '', generated_content)  # Remove markdown headers
                                generated_content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', generated_content)  # Remove markdown links
                                generated_content = re.sub(r'`([^`]+)`', r'\1', generated_content)  # Remove code formatting
                                generated_content = re.sub(r'\*\*([^*]+)\*\*', r'\1', generated_content)  # Remove bold formatting
                                generated_content = re.sub(r'\*([^*]+)\*', r'\1', generated_content)  # Remove italic formatting
                                generated_content = re.sub(r'```text\s*', '', generated_content)  # Remove code block markers
                                generated_content = re.sub(r'```\s*', '', generated_content)  # Remove remaining code block markers
                                generated_content = re.sub(r'Summary\s*', '', generated_content, flags=re.IGNORECASE)  # Remove summary labels
                                generated_content = re.sub(r'Skills\s*', '', generated_content, flags=re.IGNORECASE)  # Remove skills labels
                                generated_content = re.sub(r'Experience\s*', '', generated_content, flags=re.IGNORECASE)  # Remove experience labels
                                generated_content = re.sub(r'Education\s*', '', generated_content, flags=re.IGNORECASE)  # Remove education labels
                                generated_content = re.sub(r'Projects\s*', '', generated_content, flags=re.IGNORECASE)  # Remove projects labels
                                generated_content = re.sub(r'Languages\s*', '', generated_content, flags=re.IGNORECASE)  # Remove languages labels
                                generated_content = re.sub(r'Certifications\s*', '', generated_content, flags=re.IGNORECASE)  # Remove certifications labels
                                generated_content = re.sub(r'\n\s*\n\s*\n', '\n\n', generated_content)  # Remove excessive line breaks
                                generated_content = generated_content.strip()
                            
                            logger.info(f"Content generated successfully with {model_name} (length: {len(generated_content)} characters)")
                            break
                        else:
                            logger.warning(f"Model {model_name} returned empty response")
                            
                    except Exception as model_error:
                        logger.warning(f"Model {model_name} failed: {str(model_error)}")
                        continue
                
                if not generated_content:
                    logger.error("All Gemini models failed to generate content")
                    document.delete()
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return JsonResponse({
                            'success': False,
                            'error': 'Erreur lors de la génération du contenu. Veuillez vérifier votre clé API Gemini et réessayer.'
                        })
                    return render(request, 'user/generate.html', {
                        'error': 'Erreur lors de la génération du contenu. Veuillez vérifier votre clé API Gemini et réessayer.',
                        'templates': get_available_templates()
                    })
                    
            except Exception as e:
                logger.error(f"Gemini API error: {str(e)}")
                document.delete()
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': False,
                        'error': f'Erreur lors de la génération du document: {str(e)}. Veuillez vérifier votre connexion et réessayer.'
                    })
                return render(request, 'user/generate.html', {
                    'error': f'Erreur lors de la génération du document: {str(e)}. Veuillez vérifier votre connexion et réessayer.',
                    'templates': get_available_templates()
                })

            # Update document with enhanced scoring
            document.statut = 'completed'
            
            # Calculate dynamic score based on content quality
            score = 85  # Base score
            if len(generated_content) > 1000:
                score += 5  # Bonus for comprehensive content
            if keywords and any(keyword.lower() in generated_content.lower() for keyword in keywords.split(',')):
                score += 5  # Bonus for keyword integration
            if user_data.get('experiences') and len(user_data.get('experiences', [])) > 0:
                score += 3  # Bonus for experience data
            if user_data.get('education') and len(user_data.get('education', [])) > 0:
                score += 2  # Bonus for education data
            
            document.score = min(score, 100)  # Cap at 100
            document.contenu = generated_content
            document.save()
            logger.info(f"Document {document.id} updated to completed, score: {document.score}")

            # Mark steps as completed
            for etape in document.etape_traitement_set.all():
                etape.statut = 'completed'
                etape.save()
            logger.debug("All processing steps marked as completed")

            logger.info(f"Redirecting to dashboard for document {document.id}")
            return redirect('comptes:dashboard')

        except Exception as e:
            logger.error(f"Error in generate_document: {str(e)}", exc_info=True)
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'error': f'Une erreur s\'est produite: {str(e)}'
                })
            return render(request, 'user/generate.html', {
                'error': f'Une erreur s\'est produite: {str(e)}',
                'templates': get_available_templates()
            })

    elif request.method == 'GET':
        context = {'templates': get_available_templates()}
        doc_id = request.GET.get('doc_id')
        template_id = request.GET.get('template_id')
        if doc_id:
            document = get_object_or_404(Document, id=doc_id, user=request.user)
            context.update({
                'document': document,
                'targetRole': document.poste,
                'company': document.entreprise,
                'keywords': document.metadata.get('keywords', ''),
                'tone': document.metadata.get('tone', 'professionnel'),
                'jobDescription': document.metadata.get('job_description_preview', '')[:100],
                'documentType': document.type,
                'linkedin_url': document.linkedin_url,
                'github_url': document.github_url,
                'telephone': document.telephone,
                'langue': document.langue,
                'template_id': document.metadata.get('template_id', ''),
                'template_utilise': document.metadata.get('template_utilise', 'Template1 Moderne Bleu'),
                'skills': ', '.join(json.loads(document.metadata.get('skills', '[]'))),
                'experiences': json.loads(document.metadata.get('experiences', '[]')),
                'education': json.loads(document.metadata.get('education', '[]')),
            })
        if template_id:
            context['template_id'] = template_id
            context['template_utilise'] = next((t['name'] for t in get_available_templates() if t['id'] == template_id), 'Template1 Moderne Bleu')
        logger.info("Rendering generate.html for GET request")
        return render(request, 'user/generate.html', context)
    
    else:
        logger.warning(f"Method {request.method} not allowed for /agent/generate/")
        return HttpResponse(status=405, content="Method Not Allowed")

@csrf_exempt
@login_required
def delete_document(request, document_id):
    """API to delete a document"""
    logger.info(f"Deleting document {document_id} for user {request.user.email}")
    if request.method == 'DELETE':
        document = get_object_or_404(Document, id=document_id, user=request.user)
        try:
            document.delete()
            logger.info(f"Document {document_id} deleted successfully")
            return JsonResponse({'success': True, 'message': 'Document deleted'})
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            return JsonResponse({'success': False, 'error': str(e)})
    logger.warning(f"Method {request.method} not allowed for delete document")
    return JsonResponse({'success': False, 'error': 'Method not allowed'})

@csrf_exempt
def test_post_endpoint(request):
    """Debug endpoint to test POST requests"""
    logger.info(f"[test_post_endpoint] {time.strftime('%Y-%m-%d %H:%M:%S')} | Method: {request.method} | Path: {request.path} | Headers: {dict(request.headers)}")
    if request.method == 'POST':
        logger.debug(f"POST Data: {request.POST.dict()}")
        return JsonResponse({'success': True, 'message': 'POST request received', 'data': request.POST.dict()})
    elif request.method == 'GET':
        logger.info("GET request to test_post_endpoint")
        return JsonResponse({'success': True, 'message': 'GET request received'})
    else:
        logger.warning(f"Method {request.method} not allowed for test_post_endpoint")
        return HttpResponse(status=405, content="Method Not Allowed")

def _get_prompt(document_type, target_role, company, keywords, tone, job_description, user_data, context, langue, template_id):
    """Generate highly professional prompt for CV or LM with advanced AI instructions"""
    logger.info(f"Generating professional prompt for {document_type} with template {template_id}")
    
    # Compute skills string outside f-string
    skills_list = user_data.get('skills', [])
    skills_string = ', '.join(skills_list)
    if keywords:
        skills_string = f"{skills_string}, {keywords}" if skills_string else keywords

    # Enhanced template-specific styling instructions
    template_styles = {
        'template1-moderne-bleu': {
            'description': 'Modern professional design with blue gradient sidebar, clean typography, structured layout',
            'formatting': 'Professional CV format with clear sections and bullet points'
        },
        'template2-elegant-vert': {
            'description': 'Elegant design with green accents, sophisticated layout, professional appearance',
            'formatting': 'Elegant CV format with refined sections and professional styling'
        },
        'template3-minimaliste': {
            'description': 'Minimalist design with clean lines, focused content, modern aesthetic',
            'formatting': 'Minimalist CV format with essential information and clean structure'
        },
        'template4-corporate': {
            'description': 'Corporate design with formal layout, business-focused styling',
            'formatting': 'Corporate CV format with formal sections and professional presentation'
        },
        'template5-creatif': {
            'description': 'Creative design with innovative layout, artistic elements, modern approach',
            'formatting': 'Creative CV format with innovative sections and artistic presentation'
        },
        'template5-tech-startup': {
            'description': 'Tech startup style with dynamic layout, modern tech focus, innovative design',
            'formatting': 'Tech startup CV format with modern sections and tech-focused presentation'
        },
        'template1-classique-elegante': {
            'description': 'Classic elegant letter with formal structure, professional tone',
            'formatting': 'Classic formal letter format with proper business structure'
        },
        'template2-moderne-pro': {
            'description': 'Modern professional letter with contemporary styling, business-focused',
            'formatting': 'Modern professional letter format with contemporary business structure'
        },
        'template3-corporate-minimaliste': {
            'description': 'Corporate minimalist letter with clean design, formal business tone',
            'formatting': 'Corporate minimalist letter format with clean business structure'
        },
        'template4-creatif-colore': {
            'description': 'Creative colorful letter with innovative design, engaging presentation',
            'formatting': 'Creative letter format with innovative structure and engaging presentation'
        }
    }

    template_style = template_styles.get(template_id, template_styles['template1-moderne-bleu'])

    if document_type == 'CV':
        prompt = f"""
You are an expert CV writer and career consultant with 15+ years of experience in recruitment and HR. Create a highly professional, ATS-optimized CV that will stand out to recruiters and hiring managers.

TASK: Generate a world-class CV in {langue} for a {target_role} position at {company}.

CANDIDATE INFORMATION:
- Name: {user_data.get('name', 'Anonymous')}
- Email: {user_data.get('email', 'N/A')}
- LinkedIn: {user_data.get('linkedin_url', 'N/A')}
- GitHub: {user_data.get('github_url', 'N/A')}
- Telephone: {user_data.get('telephone', 'N/A')}
- Skills: {skills_string}
- Experiences: {json.dumps(user_data.get('experiences', []))}
- Education: {json.dumps(user_data.get('education', []))}

TARGET POSITION:
- Role: {target_role}
- Company: {company}
- Job Description: {job_description}
- Additional Context: {context}

REQUIREMENTS:
1. Use a {tone} professional tone
2. Apply '{template_style['description']}' styling
3. Ensure ATS compatibility and keyword optimization
4. Include quantifiable achievements and metrics
5. Tailor content specifically to {company} and {target_role}
6. Use action verbs and power words
7. Maintain professional formatting

OUTPUT FORMAT (clean structured content for template integration):
Generate a professional CV content that will be seamlessly integrated into the selected template design. The content should be clean, well-structured, and ready for template rendering without any markdown formatting.

PROFESSIONAL SUMMARY:
Write a compelling 150-200 word professional summary that highlights years of experience, key achievements with metrics, unique value proposition for {company}, and career objectives aligned with {target_role}. Make it natural and engaging.

TECHNICAL SKILLS:
Create a well-organized list of technical skills relevant to {target_role}. Include programming languages, frameworks, tools, technologies, methodologies, and industry-specific skills. Format as clean bullet points without asterisks, special characters, or markdown formatting.

PROFESSIONAL EXPERIENCE:
For each role, include:
- Job title, company, location, dates
- 3-4 bullet points with quantifiable achievements
- Use action verbs (Led, Developed, Implemented, Optimized)
- Include metrics and results (increased by X%, reduced by Y%, managed Z team)
- Mention technologies and tools used
- Format each experience as a clear block

EDUCATION:
Include academic qualifications: degree, institution, graduation year, relevant coursework or specializations, academic achievements or honors, and certifications and professional development.

PROJECTS:
Highlight key projects with:
- Project name and brief description
- Technologies and tools used
- Challenges solved and results achieved
- GitHub links if available
- Impact and metrics

LANGUAGES AND CERTIFICATIONS:
Include language proficiency levels, professional certifications, industry memberships, and awards and recognitions.

IMPORTANT: Do not include any markdown formatting, asterisks, backticks, or section headers. Generate clean, professional content that will be perfectly integrated into the template design.

QUALITY STANDARDS:
- Use professional terminology and industry jargon
- Include specific metrics and quantifiable results
- Ensure content is tailored to {company}'s needs
- Maintain consistency in formatting and tone
- Optimize for ATS systems with relevant keywords
- Create compelling, achievement-focused content

Generate a CV that will impress recruiters and hiring managers at {company} for the {target_role} position.
        """
    else:  # LM (Motivation Letter)
        prompt = f"""
You are an expert career consultant and professional writer specializing in motivation letters. Create a compelling, personalized motivation letter that will capture the attention of hiring managers and recruiters.

TASK: Generate a professional motivation letter in {langue} for a {target_role} position at {company}.

CANDIDATE INFORMATION:
- Name: {user_data.get('name', 'Anonymous')}
- Email: {user_data.get('email', 'N/A')}
- LinkedIn: {user_data.get('linkedin_url', 'N/A')}
- GitHub: {user_data.get('github_url', 'N/A')}
- Telephone: {user_data.get('telephone', 'N/A')}
- Skills: {skills_string}
- Experiences: {json.dumps(user_data.get('experiences', []))}
- Education: {json.dumps(user_data.get('education', []))}

TARGET POSITION:
- Role: {target_role}
- Company: {company}
- Job Description: {job_description}
- Additional Context: {context}

INSTRUCTIONS:
1. Write a natural, flowing motivation letter in {langue}
2. Use a {tone} professional tone
3. Create compelling content that connects the candidate to {company}
4. Include specific examples and achievements
5. Demonstrate knowledge of {company} and industry
6. Show enthusiasm and cultural fit
7. Write as a single, natural narrative - DO NOT include section headers or brackets

REQUIREMENTS:
- Start with a formal greeting (Madame, Monsieur,)
- Write 3-4 natural paragraphs explaining why the candidate is perfect for this role
- Include specific achievements and metrics where possible
- Show genuine interest in {company} and the {target_role} position
- End with a professional closing and contact information
- Write naturally - do not use brackets, headers, or structured format
- Make it sound like a real person wrote it, not an AI

Generate a compelling motivation letter that will make the candidate stand out and secure an interview at {company} for the {target_role} position.
        """
    
    logger.info(f"Professional prompt generated successfully for {document_type}")
    return prompt


def document_detail(request, document_id):
    """Display document details"""
    logger.info(f"Fetching document {document_id} for user {request.user.email}")
    document = get_object_or_404(Document, id=document_id, user=request.user)
    
    # Prepare context
    context = {
        'document': document,
        'etapes': document.etape_traitement_set.all().order_by('ordre')
    }
    
    # Parse content for CVs
    if document.type == 'CV':
        content = document.contenu
        # Split content into sections based on natural flow
        lines = content.split('\n')
        sections = {'profile': [], 'skills': [], 'experience': [], 'education': [], 'projects': []}
        current_section = 'profile'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect section transitions based on content patterns
            if any(word in line.lower() for word in ['compétences', 'skills', 'technologies', 'langages']):
                current_section = 'skills'
            elif any(word in line.lower() for word in ['expérience', 'experience', 'développeur', 'developer', 'ingénieur', 'engineer']):
                current_section = 'experience'
            elif any(word in line.lower() for word in ['formation', 'education', 'université', 'university', 'master', 'licence', 'diplôme']):
                current_section = 'education'
            elif any(word in line.lower() for word in ['projets', 'projects', 'portfolio', 'réalisations']):
                current_section = 'projects'
            
            sections[current_section].append(line)
        
        # Process each section's lines to identify bullet points
        def process_lines(lines):
            return [{'text': line.strip(), 'is_bullet': line.strip().startswith('- ')} for line in lines if line.strip()]
        
        context['content_sections'] = {
            'profile': {
                'title': 'Profil',
                'content': process_lines(sections['profile'] if sections['profile'] else ['No profile provided.'])
            },
            'skills': {
                'title': 'Compétences',
                'content': process_lines(sections['skills'] if sections['skills'] else [])
            },
            'experience': {
                'title': 'Expérience Professionnelle',
                'content': process_lines(sections['experience'] if sections['experience'] else [])
            },
            'education': {
                'title': 'Formation',
                'content': process_lines(sections['education'] if sections['education'] else [])
            },
            'projects': {
                'title': 'Projets',
                'content': process_lines(sections['projects'] if sections['projects'] else [])
            }
        }
    else:
        context['content_sections'] = {
            'letter': {
                'title': 'Lettre de Motivation',
                'content': [{'text': line.strip(), 'is_bullet': False} for line in document.contenu.splitlines() if line.strip()]
            }
        }
    
    # Include CV image if applicable
    if document.type == 'CV' and hasattr(document, 'cv_image'):
        context['cv_image'] = document.cv_image
        logger.debug(f"CV image found for document {document_id}")
    
    logger.info(f"Rendering document_detail for document {document_id}")
    return render(request, 'user/generate_document.html', context)
    
@login_required
def download_document(request, document_id):
    """Download generated document in selected format"""
    logger.info(f"Downloading document {document_id} for user {request.user.email}")
    document = get_object_or_404(Document, id=document_id, user=request.user)
    
    if document.statut != 'completed':
        logger.warning(f"Document {document_id} not ready for download, status: {document.statut}")
        return JsonResponse({'error': 'Document not ready'})
    
    export_format = request.GET.get('format', 'pdf')
    logger.debug(f"Generating document {document_id} in {export_format} format")
    
    filename = f"{document.type}_{document.poste.replace(' ', '_')}"
    
    # Prepare context for templates
    context = {
        'name': request.user.full_name if hasattr(request.user, 'full_name') and request.user.full_name else 'Anonymous',
        'email': request.user.email if request.user.is_authenticated else 'N/A',
        'telephone': document.telephone,
        'github_url': document.github_url,
        'github_username': document.github_url.split('/')[-1] if document.github_url else '',
        'linkedin_url': document.linkedin_url,
        'linkedin_username': document.linkedin_url.split('/')[-1] if document.linkedin_url else '',
        'target_role': document.poste,
        'company': document.entreprise,
        'company_address': document.entreprise or 'Unknown Address',
        'today': datetime.now().strftime('%d/%m/%Y'),
        'document': document,
    }

    # Parse AI-generated content (clean text without titles for CVs)
    content = document.contenu
    if document.type == 'CV':
        # Split content into sections based on natural flow
        lines = content.split('\n')
        sections = {'profile': [], 'skills': [], 'experience': [], 'education': [], 'projects': []}
        current_section = 'profile'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect section transitions based on content patterns
            if any(word in line.lower() for word in ['compétences', 'skills', 'technologies', 'langages']):
                current_section = 'skills'
            elif any(word in line.lower() for word in ['expérience', 'experience', 'développeur', 'developer', 'ingénieur', 'engineer']):
                current_section = 'experience'
            elif any(word in line.lower() for word in ['formation', 'education', 'université', 'university', 'master', 'licence', 'diplôme']):
                current_section = 'education'
            elif any(word in line.lower() for word in ['projets', 'projects', 'portfolio', 'réalisations']):
                current_section = 'projects'
            
            sections[current_section].append(line)
        
        context['profile'] = '\n'.join(sections['profile']) if sections['profile'] else 'No profile provided.'
        context['skills'] = '\n'.join(sections['skills']) if sections['skills'] else ''
        context['experience'] = '\n'.join(sections['experience']) if sections['experience'] else ''
        context['education'] = '\n'.join(sections['education']) if sections['education'] else ''
        context['projects'] = '\n'.join(sections['projects']) if sections['projects'] else ''
    else:
        context['letter_content'] = content

    # Handle CV image
    if document.type == 'CV':
        try:
            cv_image_obj = CVImage.objects.get(document=document)
            context['cv_image'] = cv_image_obj.image.url
            logger.info(f"CV image found for document {document_id}: {cv_image_obj.image.url}")
        except CVImage.DoesNotExist:
            context['cv_image'] = ''
            logger.info(f"No CV image found for document {document_id}")
    else:
        context['cv_image'] = ''

    # Get template path
    templates = get_available_templates()
    template_id = document.metadata.get('template_id', 'template1-moderne-bleu' if document.type == 'CV' else 'template1-classique-elegante')
    template = next((t for t in templates if t['id'] == template_id and t['type'] == ('CV' if document.type == 'CV' else 'LM')), None)
    if not template:
        logger.error(f"Template {template_id} not found for document {document_id}")
        return JsonResponse({'error': 'Template not found'})
    template_path = template['path']

    if export_format == 'pdf':
        try:
            html_content = render_to_string(template_path, context)
            html = HTML(string=html_content, base_url=request.build_absolute_uri('/'))
            pdf_content = html.write_pdf()
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
            logger.info(f"PDF generated for document {document_id}, filename: {filename}.pdf")
            return response
        except Exception as e:
            logger.error(f"WeasyPrint error for document {document_id}: {str(e)}")
            return JsonResponse({'error': f'PDF generation failed: {str(e)}'})
    elif export_format == 'docx':
        docx_content = generate_docx(document)
        response = HttpResponse(docx_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        logger.info(f"Word document generated for document {document_id}, filename: {filename}.docx")
        return response
    elif export_format == 'txt':
        txt_content = document.contenu.encode('utf-8')
        response = HttpResponse(txt_content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
        logger.info(f"TXT document generated for document {document_id}, filename: {filename}.txt")
        return response
    elif export_format in ['png', 'jpg']:
        html_content = render_to_string(template_path, context)
        buffer = BytesIO()
        HTML(string=html_content, base_url=request.build_absolute_uri('/')).write_png(buffer) if export_format == 'png' else HTML(string=html_content, base_url=request.build_absolute_uri('/')).write_jpg(buffer)
        image_content = buffer.getvalue()
        buffer.close()
        content_type = 'image/png' if export_format == 'png' else 'image/jpeg'
        response = HttpResponse(image_content, content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{filename}.{export_format}"'
        logger.info(f"Image ({export_format}) generated for document {document_id}, filename: {filename}.{export_format}")
        return response
    else:
        logger.warning(f"Invalid export format: {export_format}")
        return JsonResponse({'error': 'Invalid export format'})

def generate_pdf(document):
    """Generate PDF from document content using ReportLab"""
    logger.debug(f"Starting PDF generation for document {document.id}")
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    
    # Add title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, 750, document.titre)
    y = 720

    # Add content (markdown rendered as plain text)
    c.setFont("Helvetica", 12)
    for line in document.contenu.split('\n'):
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = 750
        c.drawString(100, y, line[:80])  # Truncate long lines
        y -= 15

    # Add image if available (for CV)
    if document.type == 'CV' and hasattr(document, 'cv_image') and document.cv_image.image:
        try:
            c.drawImage(document.cv_image.image.path, 400, 650, width=100, height=100)
        except Exception as e:
            logger.error(f"Error adding image to PDF for document {document.id}: {str(e)}")

    c.showPage()
    c.save()
    pdf = buffer.getvalue()
    buffer.close()
    logger.debug(f"PDF generated successfully for document {document.id}")
    return pdf

def generate_docx(document):
    """Generate Word document from document content"""
    logger.debug(f"Starting Word document generation for document {document.id}")
    doc = DocxDocument()
    doc.add_heading(document.titre, 0)
    for line in document.contenu.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    docx = buffer.getvalue()
    buffer.close()
    logger.debug(f"Word document generated successfully for document {document.id}")
    return docx

@csrf_exempt
@login_required
def upload_cv_image(request, document_id):
    """API to upload CV image"""
    logger.info(f"Uploading CV image for document {document_id} by user {request.user.email}")
    if request.method == 'POST':
        document = get_object_or_404(Document, id=document_id, user=request.user)
        
        if document.type != 'CV':
            logger.warning(f"Document {document_id} is not a CV, image upload rejected")
            return JsonResponse({'success': False, 'error': 'Document must be a CV'})
        
        if 'image' not in request.FILES:
            logger.warning("No image provided in upload request")
            return JsonResponse({'success': False, 'error': 'No image provided'})
        
        image_file = request.FILES['image']
        logger.debug(f"Image file: {image_file.name}, size: {image_file.size}")
        cv_image, created = CVImage.objects.get_or_create(document=document)
        cv_image.image = image_file
        cv_image.description = f"Professional photo for {document.poste}"
        cv_image.save()
        
        logger.info(f"Image uploaded successfully for document {document_id}")
        return JsonResponse({
            'success': True,
            'message': 'Image uploaded successfully',
            'image_url': cv_image.image.url
        })
    
    logger.warning(f"Method {request.method} not allowed for image upload")
    return JsonResponse({'success': False, 'error': 'Method not allowed'})

@login_required
def template_preview(request, template_id):
    """Preview template with sample data"""
    logger.info(f"Generating preview for template {template_id}")
    
    # Get template info
    templates = get_available_templates()
    template = next((t for t in templates if t['id'] == template_id), None)
    
    if not template:
        logger.error(f"Template {template_id} not found")
        return JsonResponse({'error': 'Template not found'})
    
    # Sample data for preview
    sample_data = {
        'name': 'Fotso Eddy Steve',
        'email': 'eddysteve@gmail.com',
        'telephone': '+1 (254) 325-2308',
        'linkedin_url': 'https://linkedin.com/in/fotsoeddysteve',
        'github_url': 'https://github.com/fotsoeddy',
        'target_role': 'Développeur Full Stack Python / JavaScript',
        'company': 'TechCorp Solutions',
        'profile': 'Développeur Full Stack passionné et orienté résultats avec plus de 7 ans d\'expérience dans la conception, le développement et le déploiement d\'applications web performantes et évolutives. Expert en Python (Django/Django REST Framework) et JavaScript (React), je maîtrise les environnements cloud et les méthodologies Agile/Scrum.',
        'skills': 'Python, JavaScript, React, Django, PostgreSQL, AWS, Docker, Git',
        'experience': '''Développeur Backend Django | HooYia Technologies | 2022-Present
- Développé des API REST sécurisées avec Django REST Framework
- Amélioré les performances de 40% et réduit les coûts d'infrastructure de 25%
- Dirigé une équipe de 3 développeurs dans la mise en œuvre de nouvelles fonctionnalités

Développeur Full Stack | ABC Company | 2020-2022
- Développé des applications web interactives avec React et Node.js
- Optimisé les performances des bases de données PostgreSQL
- Collaboré avec les équipes design et produit pour améliorer l'expérience utilisateur''',
        'education': '''Master en Informatique | Université de Technologie | 2021
- Spécialisation: Ingénierie Logicielle et Systèmes Distribués
- Mention: Bien

Licence en Informatique | Université de Sciences | 2019''',
        'projects': '''Plateforme e-commerce | React, Django, PostgreSQL
- Conception et développement d'une plateforme e-commerce complète
- Optimisation des performances pour gérer un grand nombre de produits
- Résultats: Lancement réussi avec une augmentation de 20% des ventes en ligne

Application de gestion de projet | React, Node.js, MongoDB
- Développement d'une application web pour la gestion de projets
- Implémentation d'une collaboration en temps réel
- Résultats: Amélioration de la productivité des équipes de 15%'''
    }
    
    # Add CV image if it's a CV template
    if template['type'] == 'CV':
        # Use a placeholder image data URI
        sample_data['cv_image'] = 'data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIwIiBoZWlnaHQ9IjEyMCIgdmlld0JveD0iMCAwIDEyMCAxMjAiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxyZWN0IHdpZHRoPSIxMjAiIGhlaWdodD0iMTIwIiBmaWxsPSIjRjNGNEY2Ii8+CjxjaXJjbGUgY3g9IjYwIiBjeT0iNDAiIHI9IjIwIiBmaWxsPSIjOUNBM0FGIi8+CjxwYXRoIGQ9Ik0yMCA5MEMyMCA4MC4zNTg5IDI3LjM1ODkgNzMgMzcgNzNIMTAzQzExMi42NDEgNzMgMTIwIDgwLjM1ODkgMTIwIDkwVjExMEMxMjAgMTE5LjY0MSAxMTIuNjQxIDEyNyAxMDMgMTI3SDM3QzI3LjM1ODkgMTI3IDIwIDExOS42NDEgMjAgMTEwVjkwWiIgZmlsbD0iIzlDQTNBRiIvPgo8L3N2Zz4K'
    
    # Render template with sample data
    try:
        template_path = template['path']
        html_content = render_to_string(template_path, sample_data)
        
        return JsonResponse({
            'success': True,
            'template_id': template_id,
            'template_name': template['name'],
            'template_type': template['type'],
            'preview_html': html_content
        })
    except Exception as e:
        logger.error(f"Error generating template preview: {str(e)}")
        return JsonResponse({'error': f'Preview generation failed: {str(e)}'})

@login_required
def template_selection(request):
    """Template selection page with previews"""
    logger.info(f"Template selection page accessed by user {request.user.email}")
    
    document_type = request.GET.get('document_type', 'CV')
    templates = get_available_templates()
    
    # Filter templates by document type
    filtered_templates = [t for t in templates if t['type'] == document_type]
    
    context = {
        'templates': filtered_templates,
        'document_type': document_type,
        'all_templates': templates
    }
    
    return render(request, 'user/template_selection.html', context)

@csrf_exempt
@login_required
def update_document_status(request, document_id):
    """API to update document status"""
    logger.info(f"Updating status for document {document_id} by user {request.user.email}")
    if request.method == 'POST':
        document = get_object_or_404(Document, id=document_id, user=request.user)
        try:
            data = json.loads(request.body)
            new_status = data.get('status')
            score = data.get('score', 0)
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error in update_document_status: {str(e)}")
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'})
        
        if new_status in dict(Document.STATUS_CHOICES).keys():
            document.statut = new_status
            document.score = score
            document.save()
            logger.info(f"Document {document_id} status updated to {new_status}, score: {score}")
            return JsonResponse({'success': True})
        
        logger.warning(f"Invalid status {new_status} for document {document_id}")
        return JsonResponse({'success': False, 'error': 'Invalid status'})
    
    logger.warning(f"Method {request.method} not allowed for status update")
    return JsonResponse({'success': False, 'error': 'Method not allowed'})

@login_required
def template_selection(request):
    """Display available templates for CV and LM"""
    logger.info(f"Fetching templates for user {request.user.email}")
    document_type = request.GET.get('document_type', 'CV')
    templates = [t for t in get_available_templates() if t['type'] == document_type]
    context = {
        'templates': templates,
        'document_type': document_type
    }
    return render(request, 'user/template_selection.html', context)